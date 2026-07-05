"""导出服务：排版格式化引擎 + TXT/EPUB/DOCX 导出器"""

import io
from dataclasses import dataclass
from html import escape as html_escape

from src.api.models.export_models import ExportTemplate, TemplateOptions


@dataclass
class ChapterData:
    number: int
    title: str
    content: str
    volume_number: int | None = None
    volume_title: str | None = None


PRESET_TEMPLATES: dict[ExportTemplate, TemplateOptions] = {
    ExportTemplate.default: TemplateOptions(
        chapter_title_format="第{num}章 {title}",
        paragraph_indent=2,
        paragraph_spacing=0,
        include_volume_page=True,
    ),
    ExportTemplate.qidian: TemplateOptions(
        chapter_title_format="第{num}章 {title}",
        paragraph_indent=2,
        paragraph_spacing=1,
        include_volume_page=True,
    ),
    ExportTemplate.fanqie: TemplateOptions(
        chapter_title_format="第{num}章 {title}",
        paragraph_indent=2,
        paragraph_spacing=0,
        include_volume_page=False,
    ),
}


class FormatEngine:

    def __init__(
        self, template: ExportTemplate, options: TemplateOptions | None = None
    ):
        if template == ExportTemplate.custom and options:
            self.options = options
        else:
            self.options = PRESET_TEMPLATES.get(
                template, PRESET_TEMPLATES[ExportTemplate.default]
            )

    def format_title(self, chapter: ChapterData) -> str:
        return self.options.chapter_title_format.format(
            num=chapter.number, title=chapter.title or ""
        )

    def format_paragraphs(self, content: str) -> list[str]:
        if not content:
            return []
        indent = "　" * self.options.paragraph_indent
        paragraphs = []
        for line in content.split("\n"):
            line = line.strip()
            if line:
                paragraphs.append(f"{indent}{line}")
        return paragraphs

    def format_chapter_text(self, chapter: ChapterData) -> str:
        title = self.format_title(chapter)
        paragraphs = self.format_paragraphs(chapter.content)
        spacing = "\n" * (self.options.paragraph_spacing + 1)
        body = spacing.join(paragraphs)
        return f"{title}\n\n{body}"

    def format_volume_page(self, volume_number: int, volume_title: str | None) -> str:
        title = volume_title or f"第{volume_number}卷"
        return f"{'=' * 20}\n{title}\n{'=' * 20}"


class TxtExporter:

    def __init__(self, engine: FormatEngine):
        self.engine = engine

    def export(self, chapters: list[ChapterData]) -> io.BytesIO:
        buf = io.BytesIO()

        # Prepend Table of Contents (TOC)
        if len(chapters) > 1:
            buf.write("====== 目录 ======\n\n".encode())
            current_volume: int | None = None
            for ch in chapters:
                if (
                    self.engine.options.include_volume_page
                    and ch.volume_number is not None
                    and ch.volume_number != current_volume
                ):
                    current_volume = ch.volume_number
                    vol_title = ch.volume_title or f"第{ch.volume_number}卷"
                    buf.write(f"【{vol_title}】\n".encode())

                ch_title = self.engine.format_title(ch)
                buf.write(f"  {ch_title}\n".encode())
            buf.write(b"\n==================\n\n\n")

        current_volume = None
        for ch in chapters:
            if (
                self.engine.options.include_volume_page
                and ch.volume_number is not None
                and ch.volume_number != current_volume
            ):
                current_volume = ch.volume_number
                vol_page = self.engine.format_volume_page(
                    ch.volume_number, ch.volume_title
                )
                buf.write(f"\n\n{vol_page}\n\n".encode())

            text = self.engine.format_chapter_text(ch)
            buf.write(f"{text}\n\n".encode())

        buf.seek(0)
        return buf


class EpubExporter:

    def __init__(self, engine: FormatEngine):
        self.engine = engine

    def export(
        self, chapters: list[ChapterData], title: str, author: str = "xIaoShuo"
    ) -> io.BytesIO:
        import hashlib

        from ebooklib import epub

        book = epub.EpubBook()
        uid = hashlib.md5(
            f"{title}-{len(chapters)}".encode()
        ).hexdigest()[:12]
        book.set_identifier(f"xiaoshuo-{uid}")
        book.set_title(title)
        book.set_language("zh-CN")
        book.add_author(author)

        spine: list[str | epub.EpubHtml] = ["nav"]
        toc: list[epub.EpubHtml | tuple] = []
        current_volume: int | None = None

        for ch in chapters:
            if (
                self.engine.options.include_volume_page
                and ch.volume_number is not None
                and ch.volume_number != current_volume
            ):
                current_volume = ch.volume_number
                vol_title = ch.volume_title or f"第{ch.volume_number}卷"
                vol_page = epub.EpubHtml(
                    title=vol_title,
                    file_name=f"vol_{ch.volume_number}.xhtml",
                    lang="zh-CN",
                )
                vol_page.content = (
                    f"<h1>{vol_title}</h1>".encode()
                )
                book.add_item(vol_page)
                spine.append(vol_page)

            ch_title = self.engine.format_title(ch)
            paragraphs = self.engine.format_paragraphs(ch.content)
            p_html = "".join(
                f"<p>{html_escape(p)}</p>" for p in paragraphs
            )
            html_content = (
                f"<h2>{html_escape(ch_title)}</h2>{p_html}"
            )

            epub_ch = epub.EpubHtml(
                title=ch_title,
                file_name=f"chapter_{ch.number}.xhtml",
                lang="zh-CN",
            )
            epub_ch.content = html_content.encode()
            book.add_item(epub_ch)
            spine.append(epub_ch)
            toc.append(epub_ch)

        book.toc = toc
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = spine

        buf = io.BytesIO()
        epub.write_epub(buf, book)
        buf.seek(0)
        return buf


class DocxExporter:

    def __init__(self, engine: FormatEngine):
        self.engine = engine

    def export(self, chapters: list[ChapterData], title: str) -> io.BytesIO:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(title, level=0)

        current_volume: int | None = None

        for ch in chapters:
            if (
                self.engine.options.include_volume_page
                and ch.volume_number is not None
                and ch.volume_number != current_volume
            ):
                current_volume = ch.volume_number
                doc.add_page_break()
                vol_title = ch.volume_title or f"第{ch.volume_number}卷"
                doc.add_heading(vol_title, level=1)

            ch_title = self.engine.format_title(ch)
            doc.add_heading(ch_title, level=2)

            paragraphs = self.engine.format_paragraphs(ch.content)
            for p_text in paragraphs:
                para = doc.add_paragraph(p_text)
                para.paragraph_format.first_line_indent = Pt(0)
                if self.engine.options.paragraph_spacing > 0:
                    para.paragraph_format.space_after = Pt(
                        self.engine.options.paragraph_spacing * 6
                    )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


class MobiExporter:

    def __init__(self, engine: FormatEngine):
        self.engine = engine

    def export(self, chapters: list[ChapterData], title: str) -> io.BytesIO:
        import os
        import subprocess
        import tempfile
        from shutil import which

        # Check for Calibre's ebook-convert tool
        ebook_convert = which("ebook-convert")
        if not ebook_convert:
            raise RuntimeError(
                "Calibre's 'ebook-convert' command-line tool was not found on this server. "
                "Please install Calibre to support MOBI output, or choose EPUB format instead."
            )

        # 1. Export to EPUB in memory first
        epub_buf = EpubExporter(self.engine).export(chapters, title)

        # 2. Convert to MOBI via temp files
        with tempfile.TemporaryDirectory() as tmpdir:
            epub_path = os.path.join(tmpdir, "book.epub")
            mobi_path = os.path.join(tmpdir, "book.mobi")

            with open(epub_path, "wb") as f:
                f.write(epub_buf.getvalue())

            try:
                subprocess.run(
                    [ebook_convert, epub_path, mobi_path],
                    check=True,
                    capture_output=True,
                    timeout=120,
                )

                with open(mobi_path, "rb") as f:
                    mobi_buf = io.BytesIO(f.read())
                return mobi_buf
            except Exception as e:
                raise RuntimeError(f"Failed to convert book to MOBI: {str(e)}")

