"""导出服务单元测试"""

import io

import pytest

from src.api.models.export_models import (
    ExportFormat,
    ExportRequest,
    ExportScope,
    ExportTemplate,
    TemplateOptions,
)
from src.api.services.export_service import (
    ChapterData,
    DocxExporter,
    EpubExporter,
    FormatEngine,
    TxtExporter,
)


@pytest.fixture
def sample_chapters():
    return [
        ChapterData(
            number=1,
            title="初入江湖",
            content="少年站在山巅，望着远方的云海。\n他知道，自己的命运将从此改变。",
            volume_number=1,
            volume_title="起源篇",
        ),
        ChapterData(
            number=2,
            title="奇遇",
            content="山洞深处，一道金光闪过。\n少年伸手触碰那块古老的石碑。",
            volume_number=1,
            volume_title="起源篇",
        ),
        ChapterData(
            number=3,
            title="突破",
            content="经过三天三夜的修炼，他终于突破了第一重境界。",
            volume_number=2,
            volume_title="修炼篇",
        ),
    ]


class TestFormatEngine:
    def test_default_template_title(self):
        engine = FormatEngine(ExportTemplate.default)
        ch = ChapterData(number=5, title="测试章节", content="")
        assert engine.format_title(ch) == "第5章 测试章节"

    def test_qidian_template_spacing(self):
        engine = FormatEngine(ExportTemplate.qidian)
        assert engine.options.paragraph_spacing == 1

    def test_fanqie_template_no_volume_page(self):
        engine = FormatEngine(ExportTemplate.fanqie)
        assert engine.options.include_volume_page is False

    def test_custom_template(self):
        opts = TemplateOptions(
            chapter_title_format="Chapter {num}: {title}",
            paragraph_indent=4,
            paragraph_spacing=0,
        )
        engine = FormatEngine(ExportTemplate.custom, opts)
        ch = ChapterData(number=1, title="Hello", content="")
        assert engine.format_title(ch) == "Chapter 1: Hello"

    def test_format_paragraphs_indent(self):
        engine = FormatEngine(ExportTemplate.default)
        result = engine.format_paragraphs("第一段\n第二段")
        assert len(result) == 2
        assert result[0].startswith("　　")

    def test_format_paragraphs_empty(self):
        engine = FormatEngine(ExportTemplate.default)
        assert engine.format_paragraphs("") == []
        assert engine.format_paragraphs(None) == []

    def test_format_chapter_text(self):
        engine = FormatEngine(ExportTemplate.default)
        ch = ChapterData(number=1, title="开始", content="正文内容")
        text = engine.format_chapter_text(ch)
        assert "第1章 开始" in text
        assert "　　正文内容" in text

    def test_format_volume_page(self):
        engine = FormatEngine(ExportTemplate.default)
        page = engine.format_volume_page(1, "起源篇")
        assert "起源篇" in page


class TestTxtExporter:
    def test_export_basic(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.default)
        exporter = TxtExporter(engine)
        buf = exporter.export(sample_chapters)
        text = buf.read().decode("utf-8")
        assert "第1章 初入江湖" in text
        assert "第2章 奇遇" in text
        assert "第3章 突破" in text

    def test_export_utf8(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.default)
        exporter = TxtExporter(engine)
        buf = exporter.export(sample_chapters)
        content = buf.read()
        content.decode("utf-8")

    def test_export_volume_pages(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.qidian)
        exporter = TxtExporter(engine)
        buf = exporter.export(sample_chapters)
        text = buf.read().decode("utf-8")
        assert "起源篇" in text
        assert "修炼篇" in text

    def test_export_no_volume_pages(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.fanqie)
        exporter = TxtExporter(engine)
        buf = exporter.export(sample_chapters)
        text = buf.read().decode("utf-8")
        assert "起源篇" not in text


class TestEpubExporter:
    def test_export_returns_bytes(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.default)
        exporter = EpubExporter(engine)
        buf = exporter.export(sample_chapters, title="测试小说")
        assert isinstance(buf, io.BytesIO)
        data = buf.read()
        assert len(data) > 0
        assert data[:2] == b"PK"

    def test_export_chapter_count(self, sample_chapters):
        import tempfile
        import os
        from ebooklib import epub

        engine = FormatEngine(ExportTemplate.default)
        exporter = EpubExporter(engine)
        buf = exporter.export(sample_chapters, title="测试小说")

        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
            tmp.write(buf.getvalue())
            tmp_path = tmp.name

        try:
            book = epub.read_epub(tmp_path)
            html_items = [
                i for i in book.get_items()
                if i.get_type() == 9
            ]
            assert len(html_items) >= 3
        finally:
            os.unlink(tmp_path)


class TestDocxExporter:
    def test_export_returns_bytes(self, sample_chapters):
        engine = FormatEngine(ExportTemplate.default)
        exporter = DocxExporter(engine)
        buf = exporter.export(sample_chapters, title="测试小说")
        assert isinstance(buf, io.BytesIO)
        data = buf.read()
        assert len(data) > 0
        assert data[:2] == b"PK"

    def test_export_has_headings(self, sample_chapters):
        from docx import Document

        engine = FormatEngine(ExportTemplate.default)
        exporter = DocxExporter(engine)
        buf = exporter.export(sample_chapters, title="测试小说")
        doc = Document(buf)
        all_texts = [p.text for p in doc.paragraphs]
        assert "测试小说" in all_texts
        headings = [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert any("初入江湖" in h for h in headings)


class TestExportModels:
    def test_valid_full_export(self):
        req = ExportRequest(format=ExportFormat.txt, scope=ExportScope.full)
        assert req.scope == ExportScope.full

    def test_volume_requires_number(self):
        with pytest.raises(ValueError, match="volume_number"):
            ExportRequest(format=ExportFormat.txt, scope=ExportScope.volume)

    def test_range_requires_start_end(self):
        with pytest.raises(ValueError, match="chapter_start"):
            ExportRequest(format=ExportFormat.txt, scope=ExportScope.range)

    def test_range_start_gt_end(self):
        with pytest.raises(ValueError, match="chapter_start must be"):
            ExportRequest(
                format=ExportFormat.txt,
                scope=ExportScope.range,
                chapter_start=10,
                chapter_end=5,
            )

    def test_valid_volume_export(self):
        req = ExportRequest(
            format=ExportFormat.epub,
            scope=ExportScope.volume,
            volume_number=1,
        )
        assert req.volume_number == 1
